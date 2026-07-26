import io
from pathlib import Path
from typing import Tuple, Optional

from docx import Document
from PIL import Image
from pypdf import PdfReader
import pytesseract

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg"}


def read_pdf(data: bytes) -> Tuple[str, int, list[str]]:
    warnings: list[str] = []
    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(f"\n--- Trang {idx} ---\n{text}")
    if not chunks:
        warnings.append("PDF có thể là file scan/ảnh. Hãy thử OCR hoặc upload ảnh rõ hơn.")
    return "\n".join(chunks).strip(), len(reader.pages), warnings


def read_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def read_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1258", "latin-1"):
        try:
            return data.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore").strip()


def read_image_ocr(data: bytes) -> Tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        image = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(image, lang="vie+eng")
        if not text.strip():
            warnings.append("OCR chưa đọc được chữ. Cần ảnh rõ hơn hoặc cài ngôn ngữ tiếng Việt cho Tesseract.")
        return text.strip(), warnings
    except Exception as exc:
        warnings.append(f"OCR lỗi: {exc}")
        return "", warnings


def extract_text_from_file(filename: str, data: bytes) -> Tuple[str, str, Optional[int], list[str]]:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"File {ext} chưa được hỗ trợ. Hỗ trợ: PDF, DOCX, TXT, PNG, JPG.")

    if ext == ".pdf":
        text, pages, warnings = read_pdf(data)
        return text, "PDF", pages, warnings
    if ext == ".docx":
        return read_docx(data), "DOCX", None, []
    if ext in {".txt", ".md"}:
        return read_text(data), "TEXT", None, []
    if ext in {".png", ".jpg", ".jpeg"}:
        text, warnings = read_image_ocr(data)
        return text, "IMAGE_OCR", None, warnings

    raise ValueError("Định dạng file chưa được hỗ trợ.")
